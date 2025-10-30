import os
import uuid
from pathlib import Path
from typing import Dict
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx2pdf import convert
from PyPDF2 import PdfReader, PdfWriter
import io
import pandas as pd
from templates import PROJECT_TEMPLATES
from finance import generate_financials
import google.generativeai as genai

MODEL = "gemini-2.0-flash-lite"
OUTPUT_DIR = Path("generated")
OUTPUT_DIR.mkdir(exist_ok=True)

# ✅ Predefined DPR classes
PROJECT_CLASSES = [
    "agro_processing", "textile_manufacturing", "solar_energy", "wind_energy",
    "food_packaging", "waste_management", "electric_vehicle_charging",
    "battery_manufacturing", "pharmaceuticals", "healthcare_clinic",
    "it_services", "construction_materials", "hospitality_resort",
    "education_institute", "cloud_kitchen", "boutique", "real_estate",
    "photo_studio", "bakery", "supermarket", "ice_cream", "organic_farm", "default"
]


# ---------------- CLASSIFICATION ----------------
def classify_project_with_gemini(short_description: str) -> str:
    """Use Gemini to classify a project into one of the predefined classes."""
    api_key = os.getenv("GEMINI_API_KEY")
    if not api_key:
        print("⚠️ GEMINI_API_KEY not found — using default classification.")
        return "default"

    genai.configure(api_key=api_key)
    prompt = f"""
    You are a classification assistant. Based on the short project description below,
    select the most relevant category from this list:
    {', '.join(PROJECT_CLASSES)}.

    Description: {short_description}

    Respond ONLY with the class name from the list. If none matches, say 'default'.
    """

    try:
        model = genai.GenerativeModel(MODEL)
        response = model.generate_content(prompt)
        classification = response.text.strip().lower()
        if classification not in PROJECT_CLASSES:
            classification = "default"
        return classification
    except Exception as e:
        print("classification error:", e)
        return "default"


# ---------------- STUB DATA ----------------
def fetch_external_data_stub(location: str) -> dict:
    """Static local data fetch stub."""
    return {
        "population_nearby": 120000,
        "avg_power_cost_per_kwh": 8.0,
        "land_rent_per_acre": 150000
    }


# ---------------- SECTION GENERATION ----------------
def generate_sections_with_gemini(sections: list, project_payload: dict) -> Dict[str, str]:
    """Generate content for each DPR section using Gemini API."""
    content = {}
    api_key = os.getenv("GEMINI_API_KEY")

    if not api_key:
        print("⚠️ API_KEY not found — using placeholder text.")
        for sec in sections:
            content[sec] = (
                f"[Auto-generated section: {sec}]\n\n"
                f"Project: {project_payload.get('title')}\n"
                f"Description: {project_payload.get('short_description')}"
            )
        return content

    genai.configure(api_key=api_key)

    for sec in sections:
        prompt = f"""
        You are creating content for a Detailed Project Report (DPR).

        Project title: {project_payload.get('title')}
        Short description: {project_payload.get('short_description')}
        Location: {project_payload.get('location')}

        Generate a professional section titled "{sec}" (~300-400 words).
        - Include clear **subtitles** for subtopics (use bold markers like **Subtitle:**).
        - Include bullet points where useful.
        - Avoid markdown syntax like ## or ###.
        """

        try:
            model = genai.GenerativeModel(MODEL)
            response = model.generate_content(prompt)
            text = response.text.strip()
        except Exception as e:
            print(f"❌ Gemini error for section '{sec}':", e)
            text = f"[Error generating section: {sec}]\n\n{e}"

        content[sec] = text

    return content


# ---------------- FINANCIAL PLOT ----------------
def plot_financials(df: pd.DataFrame, out_png: Path):
    """Generate financial chart image."""
    plt.clf()
    ax = df.plot(kind="line", marker="o")
    ax.set_title("Revenue & EBITDA (Projection)")
    ax.set_ylabel(df.columns[0])
    plt.tight_layout()
    plt.savefig(out_png, dpi=150)
    plt.close()


# ---------------- DOCX CREATION ----------------
def create_docx(project_payload: dict, sections: dict, df: pd.DataFrame, meta: dict, out_path: Path, chart_path: Path):
    """Create DOCX DPR with all sections, tables, and charts."""
    doc = Document()

    # Cover Page
    doc.add_heading(project_payload.get("title"), level=0)
    doc.add_paragraph(f"Location: {project_payload.get('location') or 'N/A'}")
    doc.add_paragraph(f"Currency: {meta.get('currency', 'INR')}")
    doc.add_paragraph(f"Capacity: {project_payload.get('capacity', 'N/A')}")
    doc.add_paragraph(f"Short Description: {project_payload.get('short_description', '')}")
    doc.add_page_break()

    # Sections
    for sec_title, body in sections.items():
        doc.add_heading(sec_title, level=1)
        paragraphs = body.split("\n")

        for para in paragraphs:
            para = para.strip()
            if not para:
                continue

            # Handle bold subtitles
            if para.startswith("**") and para.endswith("**"):
                run = doc.add_paragraph().add_run(para.strip("*").strip())
                run.bold = True
            elif "**" in para:
                p = doc.add_paragraph()
                parts = para.split("**")
                for i, part in enumerate(parts):
                    run = p.add_run(part)
                    if i % 2 == 1:
                        run.bold = True
            else:
                doc.add_paragraph(para)

        doc.add_page_break()

    # Financial Section
    doc.add_heading("Financial Projections (Summary)", level=1)
    t = doc.add_table(rows=1, cols=len(df.columns) + 1)
    hdr = t.rows[0].cells
    hdr[0].text = "Year"
    for i, c in enumerate(df.columns, start=1):
        hdr[i].text = str(c)

    for idx, row in df.iterrows():
        cells = t.add_row().cells
        cells[0].text = str(idx)
        for i, c in enumerate(df.columns, start=1):
            cells[i].text = f"{row[c]:,.2f}"

    if chart_path.exists():
        doc.add_paragraph()
        doc.add_picture(str(chart_path), width=Inches(6))
        doc.add_paragraph("Revenue & EBITDA Projection Chart", style="Caption")

    doc.save(out_path)


# ---------------- PDF CREATION ----------------
def create_pdf_from_docx_and_chart(docx_path: Path, chart_png: Path, out_pdf: Path):
    """
    Convert DOCX to PDF (preserving layout) and append a chart page.
    """
    temp_pdf = docx_path.with_suffix(".temp.pdf")

    try:
        # Step 1: Convert DOCX → PDF (high fidelity)
        convert(str(docx_path), str(temp_pdf))

        # Step 2: Create a one-page PDF containing chart
        chart_buffer = io.BytesIO()
        c = canvas.Canvas(chart_buffer, pagesize=A4)
        width, height = A4
        c.setFont("Helvetica-Bold", 14)
        c.drawCentredString(width / 2, height - 40, f"Financial Summary - {docx_path.stem}")

        if chart_png.exists():
            c.drawImage(str(chart_png), 40, height - 480, width=520, preserveAspectRatio=True, mask="auto")

        c.setFont("Helvetica", 10)
        c.drawString(40, 60, "Note: The chart represents financial projections and key ratios.")
        c.showPage()
        c.save()
        chart_buffer.seek(0)

        # Step 3: Merge PDFs
        writer = PdfWriter()
        doc_pdf = PdfReader(str(temp_pdf))
        for page in doc_pdf.pages:
            writer.add_page(page)

        chart_pdf = PdfReader(chart_buffer)
        for page in chart_pdf.pages:
            writer.add_page(page)

        with open(out_pdf, "wb") as f:
            writer.write(f)

        # Cleanup
        temp_pdf.unlink(missing_ok=True)

    except Exception as e:
        print("⚠️ Error generating combined PDF:", e)


# ---------------- MASTER GENERATOR ----------------
def generate_dpr_package(payload: dict) -> dict:
    """Main function: generates DPR content, DOCX, PDF, Excel."""
    uid = uuid.uuid4().hex[:8]
    project_type = classify_project_with_gemini(payload.get("short_description", ""))
    template = PROJECT_TEMPLATES.get(project_type, PROJECT_TEMPLATES["default"])
    df, meta = generate_financials(payload)

    # Generate textual DPR content
    sections = generate_sections_with_gemini(template["sections"], payload)

    # Output directory
    out_dir = OUTPUT_DIR / f"{uid}"
    out_dir.mkdir(parents=True, exist_ok=True)
    docx_path = out_dir / f"{uid}_dpr.docx"
    chart_png = out_dir / f"{uid}_finance.png"
    pdf_path = out_dir / f"{uid}_summary.pdf"
    excel_path = out_dir / f"{uid}_financials.xlsx"

    # Generate assets
    plot_financials(df, chart_png)
    df.to_excel(excel_path)
    create_docx(payload, sections, df, meta, docx_path, chart_png)
    create_pdf_from_docx_and_chart(docx_path, chart_png, pdf_path)

    return {
        "uid": uid,
        "project_type": project_type,
        "template_name": template["name"],
        "docx": str(docx_path),
        "pdf_summary": str(pdf_path),
        "chart_image": str(chart_png),
        "excel_financials": str(excel_path)
    }
